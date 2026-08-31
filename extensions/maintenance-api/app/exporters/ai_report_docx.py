from io import BytesIO
from typing import Any

from docx import Document


def _add_table(
    document: Document,
    columns: list[str],
    rows: list[list[Any]],
) -> None:
    table = document.add_table(rows=1, cols=max(1, len(columns)))
    table.style = "Table Grid"
    for index, title in enumerate(columns or ["内容"]):
        table.rows[0].cells[index].text = str(title)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values[: len(cells)]):
            cells[index].text = str(value)


def _format_source_entry(
    label: str,
    value: Any,
    fields: tuple[str, ...],
) -> str | None:
    if not isinstance(value, dict):
        return None
    details = [
        f"{field}={value[field]}"
        for field in fields
        if value.get(field) not in (None, "")
    ]
    if not details:
        return None
    return f"{label}({', '.join(details)})"


def _format_source_versions(report: dict[str, Any]) -> str:
    source_versions = report.get("source_versions")
    if not isinstance(source_versions, dict):
        return "Unavailable"

    parts: list[str] = []
    capture_mode = source_versions.get("capture_mode")
    if capture_mode not in (None, ""):
        parts.append(f"capture_mode={capture_mode}")

    definitions = (
        ("session", ("id", "version")),
        ("scenario_version", ("id", "version", "version_code")),
        (
            "calculation_run",
            (
                "id",
                "attempt_number",
                "engine_version",
                "formula_version",
                "input_snapshot_hash",
            ),
        ),
        ("review_run", ("id", "version", "rule_set_version")),
        ("inventory", ("snapshot_at",)),
    )
    for label, fields in definitions:
        entry = _format_source_entry(
            label,
            source_versions.get(label),
            fields,
        )
        if entry:
            parts.append(entry)

    return "; ".join(parts) or "Unavailable"


def _format_tenant_safe_citations(report: dict[str, Any]) -> str:
    parts: list[str] = []
    for citation in report.get("citations", []):
        if not isinstance(citation, dict):
            continue
        citation_id = citation.get("citation_id")
        if citation_id in (None, ""):
            continue

        text = f"[{citation_id}]"
        source_name = citation.get("source_name")
        if source_name not in (None, ""):
            text += f" {source_name}"
        page_number = citation.get("page_number")
        if page_number not in (None, ""):
            text += f" (page {page_number})"
        parts.append(text)

    return "; ".join(parts) or "None"


def _provenance_rows(report: dict[str, Any]) -> list[list[Any]]:
    return [
        ["Report version", report.get("version_number", "Unavailable")],
        ["Generated at", report.get("generated_at", "Unavailable")],
        ["Generation mode", report.get("generation_mode", "Unavailable")],
        ["Input digest", report.get("input_digest", "Unavailable")],
        ["Source versions / hashes", _format_source_versions(report)],
        [
            "Tenant-safe citations",
            _format_tenant_safe_citations(report),
        ],
    ]


def export_report_docx(report: dict[str, Any]) -> bytes:
    document = Document()
    document.add_heading(str(report["title"]), level=0)

    metadata = report.get("metadata", {})
    metadata_rows = [
        [key, value]
        for key, value in metadata.items()
    ]
    _add_table(
        document,
        ["字段", "内容"],
        _provenance_rows(report) + metadata_rows,
    )

    for section in report.get("sections", []):
        document.add_heading(str(section["title"]), level=1)
        document.add_paragraph(str(section.get("content", "")))
        for table_definition in section.get("tables", []):
            title = table_definition.get("title")
            if title:
                document.add_paragraph(str(title))
            _add_table(
                document,
                [
                    str(value)
                    for value in table_definition.get("columns", [])
                ],
                [
                    list(row)
                    for row in table_definition.get("rows", [])
                ],
            )
        citations = section.get("citations", [])
        if citations:
            document.add_paragraph(
                "引用："
                + "、".join(str(value) for value in citations)
            )

    document.add_heading("证据与引用", level=1)
    citation_rows = []
    for citation in report.get("citations", []):
        page = (
            ""
            if citation.get("page_number") is None
            else citation["page_number"]
        )
        citation_rows.append(
            [
                citation.get("citation_id", ""),
                citation.get("source_name", ""),
                page,
            ]
        )
    _add_table(
        document,
        ["引用编号", "来源", "页码"],
        citation_rows or [["无", "", ""]],
    )

    output = BytesIO()
    document.save(output)
    return output.getvalue()
