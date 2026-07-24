import json
from io import BytesIO

from app.exporters.ai_report_docx import export_report_docx
from app.exporters.ai_report_json import export_report_json
from app.exporters.ai_report_markdown import export_report_markdown
from docx import Document


def report() -> dict:
    return {
        "title": "维修器材需求报告",
        "metadata": {"version": "1.0", "任务": "保障演示"},
        "sections": [
            {
                "section_code": "management_summary",
                "title": "管理摘要",
                "content": "本次共识别 8 项需求。",
                "citations": ["E1"],
            },
            {
                "section_code": "calculation_results",
                "title": "需求计算结果",
                "content": "计算结果来自固定快照。",
                "tables": [
                    {
                        "title": "需求明细",
                        "columns": ["器材", "数量"],
                        "rows": [["A", "8"]],
                    }
                ],
            },
        ],
        "citations": [{"citation_id": "E1", "source_name": "手册", "page_number": 3}],
    }


def test_all_report_exports_are_machine_readable_and_docx_is_structured() -> None:
    markdown = export_report_markdown(report())
    json_text = export_report_json(report())
    docx_bytes = export_report_docx(report())

    assert "维修器材需求报告" in markdown
    assert json.loads(json_text)["title"] == "维修器材需求报告"
    assert docx_bytes[:2] == b"PK"

    document = Document(BytesIO(docx_bytes))
    paragraphs = "\n".join(row.text for row in document.paragraphs)
    assert "管理摘要" in paragraphs
    assert "需求计算结果" in paragraphs
    assert len(document.tables) >= 2
