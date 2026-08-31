from __future__ import annotations

import json
from io import BytesIO

from app.exporters.ai_report_docx import export_report_docx
from app.exporters.ai_report_json import export_report_json
from app.exporters.ai_report_markdown import export_report_markdown
from docx import Document


def _report() -> dict:
    return {
        "report_id": 7,
        "report_code": "RPT-C2B",
        "report_type": "MANAGEMENT_DECISION",
        "title": "C2B provenance export",
        "status": "DRAFT",
        "version_id": 12,
        "version_number": 2,
        "parent_version_id": 11,
        "template_version": "1.0",
        "input_digest": "a" * 64,
        "generation_mode": "RULE_FALLBACK",
        "generated_at": "2026-08-30T13:00:00+00:00",
        "source_versions": {
            "capture_mode": "AUTHORITATIVE_CREATE",
            "session": {"id": 5, "version": 3},
            "scenario_version": None,
            "calculation_run": {
                "id": 8,
                "attempt_number": 1,
                "engine_version": "1.0",
                "input_snapshot_hash": "b" * 64,
            },
            "review_run": None,
            "inventory": {
                "snapshot_at": "2026-08-30T12:00:00+00:00"
            },
        },
        "metadata": {"purpose": "c2b-export"},
        "sections": [
            {
                "section_code": "management_summary",
                "title": "管理摘要",
                "content": "确定性报告内容。",
                "citations": ["E-C2B"],
                "tables": [],
            }
        ],
        "citations": [
            {
                "citation_id": "E-C2B",
                "source_name": "tenant-safe source",
                "page_number": 3,
            }
        ],
    }


def _docx_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    chunks = [row.text for row in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def test_json_export_retains_serialized_version_provenance() -> None:
    payload = json.loads(export_report_json(_report()))

    assert payload["version_number"] == 2
    assert payload["parent_version_id"] == 11
    assert payload["input_digest"] == "a" * 64
    assert payload["generation_mode"] == "RULE_FALLBACK"
    assert payload["generated_at"]
    assert payload["source_versions"]["session"]["version"] == 3


def test_markdown_export_renders_fixed_version_provenance_block() -> None:
    markdown = export_report_markdown(_report())

    required = (
        "Report version",
        "Generated at",
        "Generation mode",
        "Input digest",
        "Source versions / hashes",
    )
    missing = [value for value in required if value not in markdown]
    assert not missing, (
        "C2B RED E02: Markdown fixed provenance block "
        f"is absent: {missing}"
    )
    assert "RULE_FALLBACK" in markdown
    assert "a" * 64 in markdown
    assert "E-C2B" in markdown


def test_docx_export_renders_fixed_version_provenance_table() -> None:
    text = _docx_text(export_report_docx(_report()))

    required = (
        "Report version",
        "Generated at",
        "Generation mode",
        "Input digest",
        "Source versions / hashes",
    )
    missing = [value for value in required if value not in text]
    assert not missing, (
        "C2B RED E03: DOCX fixed provenance table "
        f"is absent: {missing}"
    )
    assert "RULE_FALLBACK" in text
    assert "a" * 64 in text
    assert "E-C2B" in text


def test_human_exports_do_not_print_tenant_or_filesystem_metadata() -> None:
    report = _report()
    report["tenant_id"] = "tenant-c2b-secret"
    report["file_path"] = "C:/sensitive/report/path"
    report["database_url"] = "sqlite:///sensitive.db"

    markdown = export_report_markdown(report)
    docx_text = _docx_text(export_report_docx(report))

    for output in (markdown, docx_text):
        assert "tenant-c2b-secret" not in output
        assert "C:/sensitive/report/path" not in output
        assert "sqlite:///sensitive.db" not in output
